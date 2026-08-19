"""将 OCR Markdown 完全离线导出为可编辑公式的 Word 文档。"""

from __future__ import annotations

import os
import re
import uuid
import zipfile
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from urllib.parse import unquote, urlsplit


_MATH_RE = re.compile(
    r"(?<!\\)\$\$(?:(?!\$\$)[\s\S])*?(?<!\\)\$\$"
    r"|(?<!\\)\\\[(?:(?!\\\])[\s\S])*?(?<!\\)\\\]"
    r"|(?<!\\)\\\((?:(?!\\\))[\s\S])*?(?<!\\)\\\)"
    r"|(?<![\\$])\$(?!\$)(?:\\.|[^$\r\n])+?(?<!\\)\$(?!\$)"
)
_PAGE_MARKER_RE = re.compile(
    r"^[ \t]*<!--\s*第\s*(\d+)\s*页\s*-->[ \t]*$",
    re.MULTILINE,
)
_SEPARATOR_PAGE_MARKER_RE = re.compile(
    r"^[ \t]*(?:-{3,}|\*{3,}|_{3,})[ \t]*\r?\n"
    r"(?:[ \t]*\r?\n)*"
    r"(?P<marker>[ \t]*<!--\s*第\s*(?P<page>\d+)\s*页\s*-->[ \t]*)$",
    re.MULTILINE,
)
# XML 1.0 不接受的控制字符与代理码位。OCR 文本里偶尔会混入 \x0c 之类的
# 字符，直接写进 run 会让 lxml 抛错并导致整份文档导不出来。
_INVALID_XML_RE = re.compile(
    "[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff￾￿]"
)
# 这些标签的文字内容不属于正文，展开后不应出现在 Word 里。
_HTML_DROP_TAGS = frozenset(
    {"script", "style", "iframe", "object", "embed", "noscript", "template"}
)
# 结束时需要断段的块级标签。
_HTML_BLOCK_TAGS = frozenset(
    {
        "address", "article", "aside", "blockquote", "div", "dd", "dl", "dt",
        "fieldset", "figcaption", "figure", "footer", "form", "h1", "h2",
        "h3", "h4", "h5", "h6", "header", "hr", "li", "main", "nav", "ol",
        "p", "pre", "section", "ul",
    }
)
_HTML_BOLD_TAGS = frozenset({"b", "strong", "th"})
_HTML_ITALIC_TAGS = frozenset({"i", "em"})
_HTML_STRIKE_TAGS = frozenset({"s", "del", "strike"})
_HTML_CODE_TAGS = frozenset({"code", "tt", "kbd", "samp"})
_MAX_TABLE_COLUMNS = 64


def _clean_xml_text(text: str) -> str:
    """移除 XML 1.0 不允许的字符，保留制表符与换行。"""

    return _INVALID_XML_RE.sub("", text)


def _local_tag(node: Any) -> str:
    """返回小写标签名；注释、处理指令等返回空串。"""

    tag = getattr(node, "tag", None)
    if not isinstance(tag, str):
        return ""
    if "}" in tag:
        tag = tag.rsplit("}", 1)[1]
    return tag.lower()


def _element_to_html(element: Any) -> str:
    """把已解析的元素还原为 HTML 字符串，失败时返回空串。"""

    try:
        from lxml import etree

        return etree.tostring(element, encoding="unicode", method="html")
    except Exception:
        return ""


def _html_fragments(html: str) -> list[Any] | None:
    """把一段 HTML 片段解析成 (文本 / 元素) 列表，失败时返回 None。"""

    try:
        from lxml import html as lxml_html

        return list(lxml_html.fragments_fromstring(html))
    except Exception:
        return None


@dataclass
class ExportSummary:
    """一次 Word 导出的转换统计。"""

    formulas_converted: int = 0
    formulas_fallback: int = 0
    images_embedded: int = 0
    images_missing: int = 0

    @property
    def has_warnings(self) -> bool:
        return bool(self.formulas_fallback or self.images_missing)


class WordExportDependencyError(RuntimeError):
    """Word 导出依赖缺失或版本不可用。"""


@dataclass(frozen=True)
class _Formula:
    raw: str
    latex: str
    display: bool


_DEPENDENCY_MODULES = (
    "docx",
    "docx_equation",
    "latex2mathml",
    "lxml",
    "markdown_it",
    "PIL",
)


def dependencies_available() -> bool:
    """轻量探测导出依赖是否齐备（不真正导入，供界面置灰按钮用）。"""

    import importlib.util

    for name in _DEPENDENCY_MODULES:
        try:
            if importlib.util.find_spec(name) is None:
                return False
        except (ImportError, ValueError):
            return False
    return True


def _require_dependencies() -> None:
    try:
        import docx  # noqa: F401
        import docx_equation  # noqa: F401
        import latex2mathml  # noqa: F401
        import lxml  # noqa: F401
        import markdown_it  # noqa: F401
        import PIL  # noqa: F401
    except (ImportError, OSError) as exc:
        raise WordExportDependencyError(
            "Word 导出依赖不可用，请在当前 env 中安装 "
            "python-docx、docx-equation、latex2mathml、lxml、markdown-it-py 和 Pillow。"
        ) from exc


<<<<<<< Updated upstream
def _skip_balanced_group(latex: str, start: int) -> int:
    """从 start 处（必须是 '{'）扫描到配对的 '}'，返回其后的下标。"""
    depth = 0
    index = start
    while index < len(latex):
        char = latex[index]
        if char == "\\":
            # 转义字符（如 \{ \}）不参与括号配对
            index += 2
            continue
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return index + 1
        index += 1
    return -1


def _genfrac_replacement(
    left_delim: str,
    right_delim: str,
    thickness: str,
    style: str,
    numerator: str,
    denominator: str,
) -> str:
    """把 \\genfrac{ld}{rd}{th}{style}{num}{den} 重写为 latex2mathml 支持的等价形式。"""
    if thickness not in ("", "0", "0pt"):
        # 有分数横线：按样式映射为 dfrac / tfrac / frac
        if style == "0":
            fraction = rf"\dfrac{{{numerator}}}{{{denominator}}}"
        elif style == "1":
            fraction = rf"\tfrac{{{numerator}}}{{{denominator}}}"
        else:
            fraction = rf"\frac{{{numerator}}}{{{denominator}}}"
    else:
        # 无横线：\atop 直接堆叠（\binom 即空定界 + 0pt 的特例）
        fraction = "{" + numerator + r" \atop " + denominator + "}"
    if not (left_delim or right_delim):
        return fraction
    # \left< 不是合法定界，< / > 对应尖括号 \langle / \rangle
    left = r"\langle" if left_delim == "<" else (left_delim or ".")
    right = r"\rangle" if right_delim == ">" else (right_delim or ".")
    return rf"\left{left} {fraction} \right{right}"


def _preprocess_latex(latex: str) -> str:
    """把 latex2mathml 不支持的 LaTeX 命令改写为等价写法（目前仅 \\genfrac）。"""
    parts: list[str] = []
    position = 0
    index = 0
    while index < len(latex):
        if latex.startswith(r"\genfrac", index):
            cursor = index + len(r"\genfrac")
            arguments: list[str] = []
            for _ in range(6):
                while cursor < len(latex) and latex[cursor].isspace():
                    cursor += 1
                if cursor >= len(latex) or latex[cursor] != "{":
                    arguments = []
                    break
                end = _skip_balanced_group(latex, cursor)
                arguments.append(latex[cursor:end])
                cursor = end
            if len(arguments) == 6:
                parts.append(latex[position:index])
                # 分子/分母里可能再嵌套 \genfrac，递归改写后再替换
                arguments[4] = _preprocess_latex(arguments[4])
                arguments[5] = _preprocess_latex(arguments[5])
                parts.append(
                    _genfrac_replacement(*(argument[1:-1] for argument in arguments))
                )
                position = cursor
                index = cursor
                continue
        index += 1
    parts.append(latex[position:])
    return "".join(parts)


def _mathml_fenced_postprocess(math_root: Any, math_namespace: str) -> None:
    r"""把 <mo fence="true"> 定界符包装成 <mfenced>。

    latex2mathml 对 \left(...\right) 生成的是带 fence="true" 的 <mo>，
    docx-equation 不会把它们转成 OMML 的 <m:d> 定界符，导致 Word 里
    尖括号、竖线等显示异常。改成 <mfenced> 后 docx-equation 会生成
    正确的 <m:d> 节点。
    """
    from lxml import etree

    for mrow in list(math_root.iter(f"{{{math_namespace}}}mrow")):
        children = list(mrow)
        if len(children) < 2:
            continue
        first = children[0]
        last = children[-1]
        if (
            etree.QName(first).localname == "mo"
            and first.get("fence") == "true"
            and first.get("form") == "prefix"
            and etree.QName(last).localname == "mo"
            and last.get("fence") == "true"
            and last.get("form") == "postfix"
        ):
            open_text = first.text or ""
            close_text = last.text or ""
            mfenced = etree.Element(f"{{{math_namespace}}}mfenced")
            if open_text:
                mfenced.set("open", open_text)
            if close_text:
                mfenced.set("close", close_text)
            for child in children[1:-1]:
                mfenced.append(child)
            parent = mrow.getparent()
            if parent is not None:
                index = list(parent).index(mrow)
                parent.remove(mrow)
                parent.insert(index, mfenced)
=======
def _brace_group(source: str, start: int) -> tuple[str | None, int]:
    """从 source[start] 解析一个花括号组（支持嵌套）。

    返回 (组内容, 结束后的下标)；不是 '{' 或未闭合时返回 (None, 原下标)。
    """

    if start >= len(source) or source[start] != "{":
        return None, start
    depth = 0
    for index in range(start, len(source)):
        character = source[index]
        if character == "{":
            depth += 1
        elif character == "}":
            depth -= 1
            if depth == 0:
                return source[start + 1 : index], index + 1
    return None, start


def _skip_math_spaces(source: str, start: int) -> int:
    """跳过数学源码中的空白（``\\genfrac`` 各参数间允许空格）。"""

    while start < len(source) and source[start] in " \t\r\n":
        start += 1
    return start


def _genfrac_to_tex(
    left_delim: str,
    right_delim: str,
    thickness: str,
    mathstyle: str,
    numerator: str,
    denominator: str,
) -> str:
    """把 \\genfrac{ld}{rd}{th}{style}{num}{den} 重写为 latex2mathml
    支持的等价形式：

    - thickness 为 0 / 0pt / 空（无横线）→ ``{num \\atop den}`` 堆叠
    - 否则按 style 映射为 \\dfrac / \\frac / \\tfrac / \\frac
    - 定界符非空时补 ``\\left ... \\right``（"." 视为空定界符）
    - 分子分母中的嵌套 \\genfrac 会递归展开
    """

    numerator = _preprocess_genfrac(numerator)
    denominator = _preprocess_genfrac(denominator)

    no_bar = thickness.strip().lower() in {"", "0", "0pt"}
    style_commands = {
        "0": r"\dfrac",   # displaystyle
        "1": r"\frac",    # textstyle
        "2": r"\tfrac",   # scriptstyle
        "3": r"\frac",    # scriptscriptstyle（无直接等价，退化为 \frac）
    }
    if no_bar:
        core = "{" + numerator + r" \atop " + denominator + "}"
    else:
        core = (
            style_commands.get(mathstyle.strip(), r"\frac")
            + "{" + numerator + "}{" + denominator + "}"
        )

    left = left_delim.strip()
    right = right_delim.strip()
    left = "" if left in {"", "."} else left
    right = "" if right in {"", "."} else right
    if not left and not right:
        return core
    return (
        (r"\left" + left if left else r"\left.")
        + core
        + (r"\right" + right if right else r"\right.")
    )


def _preprocess_genfrac(latex: str) -> str:
    """把 latex2mathml 不支持的 ``\\genfrac{ld}{rd}{th}{style}{num}{den}``
    重写为等价的 \\frac / \\dfrac / \\tfrac 与 \\atop 堆叠，
    支持嵌套与多种定界符。"""

    parts: list[str] = []
    index = 0
    length = len(latex)
    while index < length:
        if (
            latex[index] == "\\"
            and latex.startswith("genfrac", index + 1)
            and index + 8 < length
            and not latex[index + 8].isalpha()
        ):
            cursor = _skip_math_spaces(latex, index + 8)
            groups: list[str] = []
            valid = True
            for _ in range(6):
                content, cursor = _brace_group(latex, cursor)
                if content is None:
                    valid = False
                    break
                groups.append(content)
                # \genfrac 各参数之间允许出现空格，逐个跳过
                cursor = _skip_math_spaces(latex, cursor)
            if valid:
                parts.append(_genfrac_to_tex(*groups))
                index = cursor
                continue
        parts.append(latex[index])
        index += 1
    return "".join(parts)


def _preprocess_latex(latex: str) -> str:
    """LaTeX 预处理：把 latex2mathml 不支持的语法改写为等价形式。

    当前实现 \\genfrac 重写；后续新增命令支持时在此链式追加即可。
    """

    return _preprocess_genfrac(latex)
>>>>>>> Stashed changes


def _latex_to_omml(latex: str, *, display: bool) -> Any:
    from docx_equation import mathml_to_omml
    from latex2mathml.converter import convert
    from lxml import etree

    mathml = convert(_preprocess_latex(latex.strip()))
    # latex2mathml 3.81.0 会把 aligned 的定位符输出成未转义的
    # `<mi>&</mi>`，先修成合法 XML，再转换为多行无边框矩阵。
    safe_mathml = re.sub(
        r"&(?!#(?:x[0-9A-Fa-f]+|\d+);|[A-Za-z][A-Za-z0-9]+;)",
        "&amp;",
        mathml,
    )
    math_root = etree.fromstring(safe_mathml.encode("utf-8"))
    math_namespace = "http://www.w3.org/1998/Math/MathML"

    _mathml_fenced_postprocess(math_root, math_namespace)

    has_aligned_environment = bool(
        re.search(
            r"\\begin\{(?:aligned|alignedat|align\*?|split|gathered)\}",
            latex,
        )
    )
    if has_aligned_environment:
        line_breaks = math_root.xpath(
            ".//*[local-name()='mspace' and @linebreak='newline']"
        )
        parents = []
        for line_break in line_breaks:
            parent = line_break.getparent()
            if parent is not None and all(parent is not item for item in parents):
                parents.append(parent)
        for parent in parents:
            rows: list[list[Any]] = [[]]
            for child in list(parent):
                parent.remove(child)
                local_name = etree.QName(child).localname
                if (
                    local_name == "mspace"
                    and child.get("linebreak") == "newline"
                ):
                    rows.append([])
                elif local_name == "mi" and (child.text or "") == "&":
                    continue
                else:
                    rows[-1].append(child)
            table = etree.Element(f"{{{math_namespace}}}mtable")
            for row_items in rows:
                if not row_items:
                    continue
                table_row = etree.SubElement(
                    table,
                    f"{{{math_namespace}}}mtr",
                )
                table_cell = etree.SubElement(
                    table_row,
                    f"{{{math_namespace}}}mtd",
                )
                row = etree.SubElement(
                    table_cell,
                    f"{{{math_namespace}}}mrow",
                )
                row.extend(row_items)
            parent.append(table)
        for alignment_marker in math_root.xpath(
            ".//*[local-name()='mi' and text()='&']"
        ):
            marker_parent = alignment_marker.getparent()
            if marker_parent is not None:
                marker_parent.remove(alignment_marker)

    sentinel = f"DOCXEQSQRT{uuid.uuid4().hex}"
    has_square_root = False

    # docx-equation 0.3.0 的 msqrt 分支会把节点自身再次递归转换。
    # 先把它规范成带临时根指数的 mroot，转换后再隐藏根指数。
    for square_root in math_root.xpath(".//*[local-name()='msqrt']"):
        has_square_root = True
        radicand = etree.Element(f"{{{math_namespace}}}mrow")
        radicand.text = square_root.text
        square_root.text = None
        for child in list(square_root):
            square_root.remove(child)
            radicand.append(child)
        degree = etree.Element(f"{{{math_namespace}}}mn")
        degree.text = sentinel
        square_root.tag = f"{{{math_namespace}}}mroot"
        square_root.extend((radicand, degree))

    omml = mathml_to_omml(
        etree.tostring(math_root, encoding="unicode"),
        display=display,
    )
    omml_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    if has_square_root:
        for radical in omml.xpath(
            ".//m:rad[m:deg//*[contains(text(), $sentinel)]]",
            namespaces={"m": omml_namespace},
            sentinel=sentinel,
        ):
            degree = radical.find(f"{{{omml_namespace}}}deg")
            if degree is not None:
                degree.clear()
            properties = radical.find(f"{{{omml_namespace}}}radPr")
            if properties is not None:
                degree_hidden = etree.Element(
                    f"{{{omml_namespace}}}degHide"
                )
                degree_hidden.set(f"{{{omml_namespace}}}val", "1")
                properties.insert(0, degree_hidden)

    word_namespace = (
        "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    )
    for fonts in omml.xpath(
        ".//w:rFonts",
        namespaces={"w": word_namespace},
    ):
        for attribute in ("ascii", "hAnsi", "cs"):
            fonts.set(
                f"{{{word_namespace}}}{attribute}",
                "Cambria Math",
            )
        fonts.set(f"{{{word_namespace}}}eastAsia", "微软雅黑")

    # docx-equation 生成的 <m:rPr>/<m:sty> 与 <m:ctrlPr> 会让 Microsoft Word
    # 严格校验报错；LibreOffice / Word 均不需要这些元素即可正确渲染。
    # 同时给 <m:t> 加 xml:space="preserve"，避免首尾空格被吃掉。
    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
    for rpr in list(omml.xpath(".//m:rPr", namespaces={"m": omml_namespace})):
        parent = rpr.getparent()
        if parent is not None:
            parent.remove(rpr)
    for ctrl_pr in list(omml.xpath(".//m:ctrlPr", namespaces={"m": omml_namespace})):
        parent = ctrl_pr.getparent()
        if parent is not None:
            parent.remove(ctrl_pr)
    for t_el in omml.xpath(".//m:t", namespaces={"m": omml_namespace}):
        text = t_el.text or ""
        if text.startswith(" ") or text.endswith(" "):
            t_el.set(XML_SPACE, "preserve")
    return omml


def _formula_details(raw: str) -> tuple[str, bool]:
    if raw.startswith("$$") and raw.endswith("$$"):
        return raw[2:-2], True
    if raw.startswith(r"\[") and raw.endswith(r"\]"):
        return raw[2:-2], True
    if raw.startswith(r"\(") and raw.endswith(r"\)"):
        return raw[2:-2], False
    return raw[1:-1], False


def _set_east_asia_font(style, name: str) -> None:
    from docx.oxml.ns import qn

    style.font.name = name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _new_document():
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Mm, Pt

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

    normal = document.styles["Normal"]
    _set_east_asia_font(normal, "微软雅黑")
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    return document


def _wrap_omml_in_run(omml) -> Any:
    """把 <m:oMath> / <m:oMathPara> 包裹到 <w:r> 中，符合 Word 的段落内容模型。"""
    from lxml import etree

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    run = etree.Element(f"{{{W_NS}}}r")
    run.append(omml)
    return run


def _append_formula(paragraph, raw: str, summary: ExportSummary) -> None:
    latex, _display = _formula_details(raw)
    try:
        paragraph._p.append(_wrap_omml_in_run(_latex_to_omml(latex, display=False)))
    except Exception:
        paragraph.add_run(_clean_xml_text(raw))
        summary.formulas_fallback += 1
    else:
        summary.formulas_converted += 1


def _append_display_formula(document, raw: str, summary: ExportSummary) -> None:
    from lxml import etree

    word_namespace = (
        "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    )

    latex, _display = _formula_details(raw)
    try:
        omml = _latex_to_omml(latex, display=True)
    except Exception:
        paragraph = document.add_paragraph(_clean_xml_text(raw))
        paragraph.alignment = 1
        summary.formulas_fallback += 1
        return

    math_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    properties = omml.find(f"{{{math_namespace}}}oMathParaPr")
    if properties is None:
        properties = etree.Element(f"{{{math_namespace}}}oMathParaPr")
        omml.insert(0, properties)
    justification = properties.find(f"{{{math_namespace}}}jc")
    if justification is None:
        justification = etree.SubElement(
            properties,
            f"{{{math_namespace}}}jc",
        )
    justification.set(f"{{{math_namespace}}}val", "center")

<<<<<<< Updated upstream
    body = document._element.body
    # OMML 规范：<m:oMathPara> 必须直接作为 <w:p> 的子元素，不能包在 <w:r> 里。
    # 先创建段落再把 oMathPara 放入，按调用顺序追加到 body 末尾。
    paragraph = etree.Element(f"{{{word_namespace}}}p")
    paragraph.append(omml)
    body.append(paragraph)
=======
    # Word 要求 <m:oMathPara> 必须位于 <w:p> 段落内，不能直接挂在 body 下。
    paragraph = document.add_paragraph()
    paragraph._p.append(omml)
>>>>>>> Stashed changes
    summary.formulas_converted += 1


def _protect_markdown_code(markdown: str) -> tuple[str, dict[str, str]]:
    prefix = f"DOCXCODE{uuid.uuid4().hex.upper()}X"
    protected_regions: dict[str, str] = {}
    lines = markdown.splitlines(keepends=True)
    line_offsets = [0]
    for line in lines:
        line_offsets.append(line_offsets[-1] + len(line))

    ranges = []
    for token in _markdown_parser().parse(markdown):
        if token.type not in {"fence", "code_block"} or token.map is None:
            continue
        start_line, end_line = token.map
        ranges.append((line_offsets[start_line], line_offsets[end_line]))

    parts: list[str] = []
    position = 0
    for start, end in sorted(ranges):
        if start < position:
            continue
        parts.append(markdown[position:start])
        placeholder = f"{prefix}{len(protected_regions)}Z"
        protected_regions[placeholder] = markdown[start:end]
        parts.append(placeholder)
        position = end
    parts.append(markdown[position:])
    block_protected = "".join(parts)

    # CommonMark 的行内代码可以跨软换行，但 OCR 文本里不成对的孤立反引号常见：
    # 若允许跨行匹配，两个孤立反引号会把中间整段（含公式）都当成代码保护起来，
    # 导致其中的公式再也不会被转换成 OMML。因此这里收紧为仅同一行内配对。
    inline_pattern = re.compile(
        r"(?<!\\)(?P<ticks>`+)(?P<body>[^\r\n]*?)(?P=ticks)"
    )

    def replace_inline(match: re.Match[str]) -> str:
        placeholder = f"{prefix}{len(protected_regions)}Z"
        protected_regions[placeholder] = match.group(0)
        return placeholder

    return inline_pattern.sub(replace_inline, block_protected), protected_regions


def _restore_markdown_code(
    markdown: str,
    protected_regions: dict[str, str],
) -> str:
    for placeholder, original in protected_regions.items():
        markdown = markdown.replace(placeholder, original)
    return markdown


def _protect_formulas(markdown: str) -> tuple[str, dict[str, _Formula]]:
    prefix = f"DOCXFORMULA{uuid.uuid4().hex.upper()}X"
    formulas: dict[str, _Formula] = {}

    def replace(match: re.Match[str]) -> str:
        raw = match.group(0)
        latex, display = _formula_details(raw)
        placeholder = f"{prefix}{len(formulas)}Z"
        formulas[placeholder] = _Formula(raw, latex, display)
        return placeholder

    return _MATH_RE.sub(replace, markdown), formulas


def _protect_page_markers(markdown: str) -> tuple[str, dict[str, int]]:
    prefix = f"DOCXPAGE{uuid.uuid4().hex.upper()}X"
    markers: dict[str, int] = {}

    def remove_separator(match: re.Match[str]) -> str:
        if int(match.group("page")) <= 1:
            return match.group(0)
        return match.group("marker")

    without_page_separators = _SEPARATOR_PAGE_MARKER_RE.sub(
        remove_separator,
        markdown,
    )

    def replace_marker(match: re.Match[str]) -> str:
        page_number = int(match.group(1))
        if page_number <= 1:
            return ""
        placeholder = f"{prefix}{len(markers)}Z"
        markers[placeholder] = page_number
        return f"\n\n{placeholder}\n\n"

    return _PAGE_MARKER_RE.sub(replace_marker, without_page_separators), markers


def _markdown_parser():
    from markdown_it import MarkdownIt

    return (
        MarkdownIt(
            "commonmark",
            {"html": True, "linkify": False, "typographer": False},
        )
        .enable("table")
        .enable("strikethrough")
    )


class _VisibleHtmlParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._blocked_depth = 0

    def handle_starttag(self, tag: str, attrs: list) -> None:
        if tag.lower() in _HTML_DROP_TAGS:
            self._blocked_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in _HTML_DROP_TAGS and self._blocked_depth:
            self._blocked_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._blocked_depth:
            self.parts.append(data)

    def visible_text(self) -> str:
        return "".join(self.parts).strip()


def _visible_html(html: str) -> str:
    parser = _VisibleHtmlParser()
    try:
        parser.feed(html)
        parser.close()
    except Exception:
        return html
    return parser.visible_text()


def _set_run_font(run, name: str) -> None:
    from docx.oxml.ns import qn

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _shade_paragraph(paragraph, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _add_horizontal_rule(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7C3D0")
    borders.append(bottom)
    properties.append(borders)


def _safe_hyperlink_target(href: str) -> str | None:
    if not href or any(ord(character) < 32 for character in href):
        return None
    try:
        parsed = urlsplit(href)
    except ValueError:
        return None
    scheme = parsed.scheme.lower()
    if scheme not in {"http", "https", "mailto"}:
        return None
    if scheme in {"http", "https"} and not parsed.netloc:
        return None
    if scheme == "mailto" and not parsed.path:
        return None
    return href


def _resolve_local_image(src: str, base_dir: Path) -> Path | None:
    from PIL import Image

    try:
        parsed = urlsplit(src)
        if parsed.scheme or parsed.netloc or not parsed.path:
            return None
        relative = Path(unquote(parsed.path))
        root = base_dir.resolve()
        candidate = (root / relative).resolve()
        if not candidate.is_relative_to(root) or not candidate.is_file():
            return None
        with Image.open(candidate) as image:
            image.verify()
        return candidate
    except (OSError, ValueError, SyntaxError):
        return None


def _add_hyperlink(paragraph, text: str, href: str, formatting: dict[str, bool]):
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(
        href,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = paragraph.add_run(text)
    run.bold = formatting["bold"]
    run.italic = formatting["italic"]
    run.font.strike = formatting["strike"]
    if formatting["code"]:
        _set_run_font(run, "Consolas")
    run_properties = run._element.get_or_add_rPr()
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    paragraph._p.remove(run._r)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    return run


class _DocxRenderer:
    def __init__(
        self,
        document,
        *,
        base_dir: Path,
        formulas: dict[str, _Formula],
        page_markers: dict[str, int],
        summary: ExportSummary,
    ) -> None:
        self.document = document
        self.base_dir = base_dir
        self.formulas = formulas
        self.page_markers = page_markers
        self.summary = summary
        self.list_stack: list[str] = []
        self.blockquote_depth = 0
        self._formula_pattern = (
            re.compile(
                "|".join(
                    re.escape(item)
                    for item in sorted(
                        formulas,
                        key=len,
                        reverse=True,
                    )
                )
            )
            if formulas
            else None
        )

    def _new_context_paragraph(self, style: str | None = None):
        from docx.shared import Mm

        selected_style = style
        if selected_style is None and self.list_stack:
            level = min(len(self.list_stack), 3)
            base = (
                "List Bullet"
                if self.list_stack[-1] == "bullet"
                else "List Number"
            )
            selected_style = base if level == 1 else f"{base} {level}"
        paragraph = self.document.add_paragraph(style=selected_style)
        if self.blockquote_depth:
            paragraph.paragraph_format.left_indent = Mm(
                7 * self.blockquote_depth
            )
            paragraph.paragraph_format.right_indent = Mm(3)
        return paragraph

    def _single_display_formula(self, inline_token) -> _Formula | None:
        children = inline_token.children or []
        if len(children) != 1 or children[0].type != "text":
            return None
        formula = self.formulas.get(children[0].content)
        if formula is None or not formula.display:
            return None
        return formula

    def _single_page_marker(self, inline_token) -> int | None:
        children = inline_token.children or []
        if len(children) != 1 or children[0].type != "text":
            return None
        return self.page_markers.get(children[0].content)

    @staticmethod
    def _formatting() -> dict[str, bool]:
        return {
            "bold": False,
            "italic": False,
            "strike": False,
            "code": False,
        }

    def _add_text(
        self,
        paragraph,
        text: str,
        formatting: dict[str, bool],
        *,
        href: str | None = None,
    ) -> None:
        text = _clean_xml_text(text)
        if not text:
            return
        if href is not None:
            _add_hyperlink(paragraph, text, href, formatting)
            return
        run = paragraph.add_run(text)
        run.bold = formatting["bold"]
        run.italic = formatting["italic"]
        run.font.strike = formatting["strike"]
        if formatting["code"]:
            _set_run_font(run, "Consolas")

    def _add_text_and_formulas(
        self,
        paragraph,
        text: str,
        formatting: dict[str, bool],
        *,
        href: str | None = None,
    ) -> None:
        if self._formula_pattern is None:
            self._add_text(paragraph, text, formatting, href=href)
            return
        position = 0
        for match in self._formula_pattern.finditer(text):
            if match.start() > position:
                self._add_text(
                    paragraph,
                    text[position : match.start()],
                    formatting,
                    href=href,
                )
            formula = self.formulas[match.group(0)]
            _append_formula(paragraph, formula.raw, self.summary)
            position = match.end()
        if position < len(text):
            self._add_text(
                paragraph,
                text[position:],
                formatting,
                href=href,
            )

    def _add_inline_html(
        self,
        paragraph,
        html: str,
        formatting: dict[str, bool],
        *,
        href: str | None = None,
    ) -> None:
        """行内 HTML：``<img>`` 走图片嵌入，其余只保留可见文字。"""
        if "<img" in html.lower():
            nodes = _html_fragments(html)
            if nodes is not None:
                handled = False
                for node in nodes:
                    if isinstance(node, str):
                        self._add_text_and_formulas(
                            paragraph, node, formatting, href=href
                        )
                        continue
                    if _local_tag(node) == "img":
                        self._add_image(paragraph, node.get("src") or "")
                        handled = True
                    else:
                        self._render_html_inline(
                            paragraph, node, force_bold=formatting["bold"]
                        )
                        handled = True
                    if node.tail:
                        self._add_text_and_formulas(
                            paragraph, node.tail, formatting, href=href
                        )
                if handled:
                    return
        self._add_text_and_formulas(
            paragraph,
            _visible_html(html),
            formatting,
            href=href,
        )

    def render_inline(
        self,
        paragraph,
        children,
        *,
        force_bold: bool = False,
    ) -> None:
        formatting = self._formatting()
        formatting["bold"] = force_bold
        link_target: str | None = None
        for token in children or []:
            token_type = token.type
            if token_type == "text":
                self._add_text_and_formulas(
                    paragraph,
                    token.content,
                    formatting,
                    href=link_target,
                )
            elif token_type == "strong_open":
                formatting["bold"] = True
            elif token_type == "strong_close":
                formatting["bold"] = force_bold
            elif token_type == "em_open":
                formatting["italic"] = True
            elif token_type == "em_close":
                formatting["italic"] = False
            elif token_type == "s_open":
                formatting["strike"] = True
            elif token_type == "s_close":
                formatting["strike"] = False
            elif token_type == "code_inline":
                code_formatting = dict(formatting)
                code_formatting["code"] = True
                self._add_text(
                    paragraph,
                    token.content,
                    code_formatting,
                    href=link_target,
                )
            elif token_type == "softbreak":
                self._add_text(paragraph, " ", formatting, href=link_target)
            elif token_type == "hardbreak":
                paragraph.add_run().add_break()
            elif token_type == "link_open":
                link_target = _safe_hyperlink_target(
                    token.attrGet("href") or ""
                )
            elif token_type == "link_close":
                link_target = None
            elif token_type == "image":
                self._render_image(paragraph, token)
            elif token_type == "html_inline":
                self._add_inline_html(
                    paragraph,
                    token.content,
                    formatting,
                    href=link_target,
                )
            elif token.content:
                self._add_text_and_formulas(
                    paragraph,
                    token.content,
                    formatting,
                    href=link_target,
                )

    def _render_image(self, paragraph, token) -> None:
        self._add_image(paragraph, token.attrGet("src") or "")

    def _add_image(self, paragraph, source: str) -> None:
        from docx.shared import Emu
        from PIL import Image

        image_path = _resolve_local_image(source, self.base_dir)
        if image_path is None:
            paragraph.add_run(_clean_xml_text(f"[图片不可用：{source}]"))
            self.summary.images_missing += 1
            return

        try:
            with Image.open(image_path) as image:
                width_pixels, _height_pixels = image.size
                # 部分 TIFF 无分辨率单位时 info 里无 dpi 键；带单位时返回 (dpi_x, dpi_y)
                dpi_value = image.info.get("dpi") or (96, 96)
            if isinstance(dpi_value, (int, float)):
                dpi_x = float(dpi_value)
            else:
                dpi_x = float(dpi_value[0])
            if not 10 <= dpi_x <= 1200:
                dpi_x = 96.0
            natural_width = int(width_pixels / dpi_x * 914400)
            section = self.document.sections[-1]
            available_width = int(
                section.page_width
                - section.left_margin
                - section.right_margin
            )
            display_width = max(1, min(natural_width, available_width))
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Emu(display_width))
        except Exception:
            paragraph.add_run(_clean_xml_text(f"[图片不可用：{source}]"))
            self.summary.images_missing += 1
            return
        self.summary.images_embedded += 1

    def _render_code(self, content: str) -> None:
        from docx.shared import Pt

        paragraph = self._new_context_paragraph()
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        _shade_paragraph(paragraph, "F3F5F7")
        run = paragraph.add_run(_clean_xml_text(content.rstrip("\r\n")))
        _set_run_font(run, "Consolas")
        run.font.size = Pt(9)

    def _render_table(self, tokens, start: int) -> int:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Pt, Twips

        rows: list[list[tuple[Any, bool]]] = []
        row: list[tuple[Any, bool]] | None = None
        cell_header = False
        index = start + 1
        while index < len(tokens) and tokens[index].type != "table_close":
            token = tokens[index]
            if token.type == "tr_open":
                row = []
            elif token.type == "tr_close" and row is not None:
                rows.append(row)
                row = None
            elif token.type in {"th_open", "td_open"}:
                cell_header = token.type == "th_open"
            elif token.type == "inline" and row is not None:
                row.append((token, cell_header))
            index += 1

        if not rows:
            return index
        column_count = min(
            max((len(item) for item in rows), default=0),
            _MAX_TABLE_COLUMNS,
        )
        if column_count < 1:
            return index
        table = self.document.add_table(
            rows=len(rows),
            cols=column_count,
        )
        column_widths = self._apply_table_layout(table, column_count)
        for row_index, cells in enumerate(rows):
            for column_index, (inline_token, is_header) in enumerate(cells):
                if column_index >= column_count:
                    break
                cell = table.cell(row_index, column_index)
                cell.width = Twips(column_widths[column_index])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                self.render_inline(
                    paragraph,
                    inline_token.children,
                    force_bold=is_header,
                )
                if is_header:
                    self._shade_cell(cell)
            if row_index == 0:
                self._repeat_header_row(table)
        return index

    def _apply_table_layout(self, table, column_count: int) -> list[int]:
        """统一设置表格样式、总宽度、单元格边距与列宽。"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        table.style = "Table Grid"
        table.autofit = False
        section = self.document.sections[-1]
        table_width = int(
            (
                section.page_width
                - section.left_margin
                - section.right_margin
            )
            / 635
        )
        base_column_width, remainder = divmod(table_width, column_count)
        column_widths = [
            base_column_width + (1 if index < remainder else 0)
            for index in range(column_count)
        ]
        table_properties = table._tbl.tblPr
        width_element = table_properties.first_child_found_in("w:tblW")
        if width_element is None:
            width_element = OxmlElement("w:tblW")
            table_properties.insert(0, width_element)
        width_element.set(qn("w:type"), "dxa")
        width_element.set(qn("w:w"), str(table_width))

        cell_margins = OxmlElement("w:tblCellMar")
        for side, value in (
            ("top", 80),
            ("start", 120),
            ("bottom", 80),
            ("end", 120),
        ):
            margin = OxmlElement(f"w:{side}")
            margin.set(qn("w:w"), str(value))
            margin.set(qn("w:type"), "dxa")
            cell_margins.append(margin)
        table_properties.append(cell_margins)

        for grid_column, width in zip(
            table._tbl.tblGrid.gridCol_lst,
            column_widths,
        ):
            grid_column.set(qn("w:w"), str(width))
        return column_widths

    @staticmethod
    def _shade_cell(cell) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E9EFF5")
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _repeat_header_row(table) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        row_properties = table.rows[0]._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        row_properties.append(repeat)

    # ------------------------------------------------------- 原始 HTML --

    @staticmethod
    def _own_table_rows(table_element) -> list[Any]:
        """只取属于该表格自身的 tr，忽略嵌套表格里的行。"""
        rows = []
        for row in table_element.iter():
            if _local_tag(row) != "tr":
                continue
            ancestor = row.getparent()
            while ancestor is not None and _local_tag(ancestor) != "table":
                ancestor = ancestor.getparent()
            if ancestor is table_element:
                rows.append(row)
        return rows

    @classmethod
    def _html_table_grid(cls, table_element):
        """把 HTML 表格解析成带合并信息的网格。"""
        rows = cls._own_table_rows(table_element)
        if not rows:
            return [], 0, 0

        def span(cell, name: str, limit: int) -> int:
            try:
                value = int(str(cell.get(name, "1")).strip() or "1")
            except (TypeError, ValueError):
                return 1
            return max(1, min(value, limit))

        occupied: set[tuple[int, int]] = set()
        placements = []
        for row_index, row in enumerate(rows):
            column_index = 0
            for cell in row:
                tag = _local_tag(cell)
                if tag not in ("td", "th"):
                    continue
                while (row_index, column_index) in occupied:
                    column_index += 1
                if column_index >= _MAX_TABLE_COLUMNS:
                    break
                # 行合并不能越过表格末尾，列合并按上限收敛。
                row_span = span(cell, "rowspan", len(rows) - row_index)
                column_span = span(
                    cell,
                    "colspan",
                    _MAX_TABLE_COLUMNS - column_index,
                )
                for delta_row in range(row_span):
                    for delta_column in range(column_span):
                        occupied.add(
                            (row_index + delta_row, column_index + delta_column)
                        )
                placements.append(
                    (row_index, column_index, row_span, column_span, cell, tag == "th")
                )
                column_index += column_span

        if not placements:
            return [], 0, 0
        total_rows = max(row + span_rows for row, _, span_rows, _, _, _ in placements)
        total_columns = min(
            max(column + span_columns for _, column, _, span_columns, _, _ in placements),
            _MAX_TABLE_COLUMNS,
        )
        return placements, total_rows, total_columns

    def _render_html_table(self, table_element) -> bool:
        """把 PaddleOCR 输出的 HTML 表格还原成真正的 Word 表格。"""
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Pt, Twips

        placements, total_rows, total_columns = self._html_table_grid(table_element)
        if not placements or total_rows < 1 or total_columns < 1:
            return False

        table = self.document.add_table(rows=total_rows, cols=total_columns)
        column_widths = self._apply_table_layout(table, total_columns)
        header_rows = {
            row for row, _, _, _, _, is_header in placements if is_header
        }
        for row, column, row_span, column_span, element, is_header in placements:
            if column >= total_columns:
                continue
            column_span = min(column_span, total_columns - column)
            cell = table.cell(row, column)
            if row_span > 1 or column_span > 1:
                try:
                    cell = cell.merge(
                        table.cell(row + row_span - 1, column + column_span - 1)
                    )
                except (IndexError, ValueError):
                    cell = table.cell(row, column)
            else:
                cell.width = Twips(column_widths[column])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            self._render_html_inline(paragraph, element, force_bold=is_header)
            if is_header:
                self._shade_cell(cell)
        if 0 in header_rows:
            self._repeat_header_row(table)
        return True

    def _render_html_inline(self, paragraph, element, *, force_bold: bool) -> None:
        """把一个 HTML 元素的内容渲染进已有段落（文本 / 公式 / 图片）。"""
        base_formatting = self._formatting()
        base_formatting["bold"] = force_bold

        def walk(node, formatting: dict[str, bool]) -> None:
            tag = _local_tag(node)
            if tag in _HTML_DROP_TAGS:
                return
            if tag == "img":
                self._add_image(paragraph, node.get("src") or "")
                return
            if tag == "br":
                paragraph.add_run().add_break()
                return
            child_formatting = dict(formatting)
            if tag in _HTML_BOLD_TAGS:
                child_formatting["bold"] = True
            if tag in _HTML_ITALIC_TAGS:
                child_formatting["italic"] = True
            if tag in _HTML_STRIKE_TAGS:
                child_formatting["strike"] = True
            if tag in _HTML_CODE_TAGS:
                child_formatting["code"] = True
            if node.text:
                self._add_text_and_formulas(
                    paragraph, node.text, child_formatting
                )
            for child in node:
                walk(child, child_formatting)
                if child.tail:
                    self._add_text_and_formulas(
                        paragraph, child.tail, child_formatting
                    )

        walk(element, base_formatting)

    def _render_html_fragment(self, html: str) -> None:
        """渲染原始 HTML 块：表格建表、图片嵌入、其余按段落输出文字。"""
        nodes = _html_fragments(html)
        if nodes is None:
            visible = _visible_html(html)
            if visible:
                self._add_text_and_formulas(
                    self._new_context_paragraph(), visible, self._formatting()
                )
            return

        pending: list[tuple[str, Any]] = []

        def flush() -> None:
            if not pending:
                return
            items = list(pending)
            pending.clear()
            if all(
                kind == "text" and not value.strip() for kind, value in items
            ):
                return
            paragraph = self._new_context_paragraph()
            for kind, value in items:
                if kind == "text":
                    self._add_text_and_formulas(
                        paragraph, value, self._formatting()
                    )
                elif kind == "image":
                    self._add_image(paragraph, value)
                else:
                    paragraph.add_run().add_break()

        def walk(node) -> None:
            if isinstance(node, str):
                pending.append(("text", node))
                return
            tag = _local_tag(node)
            if not tag:  # 注释、处理指令
                if node.tail:
                    pending.append(("text", node.tail))
                return
            if tag in _HTML_DROP_TAGS:
                if node.tail:
                    pending.append(("text", node.tail))
                return
            if tag == "table":
                flush()
                if not self._render_html_table(node):
                    visible = _visible_html(
                        _element_to_html(node) or ""
                    )
                    if visible:
                        pending.append(("text", visible))
                if node.tail:
                    pending.append(("text", node.tail))
                return
            if tag == "img":
                pending.append(("image", node.get("src") or ""))
                if node.tail:
                    pending.append(("text", node.tail))
                return
            if tag == "br":
                pending.append(("break", None))
                if node.tail:
                    pending.append(("text", node.tail))
                return
            if node.text:
                pending.append(("text", node.text))
            for child in node:
                walk(child)
            if tag in _HTML_BLOCK_TAGS:
                flush()
            if node.tail:
                pending.append(("text", node.tail))

        for node in nodes:
            walk(node)
        flush()

    def render(self, tokens) -> None:
        index = 0
        while index < len(tokens):
            token = tokens[index]
            token_type = token.type
            if token_type == "bullet_list_open":
                self.list_stack.append("bullet")
            elif token_type == "ordered_list_open":
                self.list_stack.append("number")
            elif token_type in {"bullet_list_close", "ordered_list_close"}:
                if self.list_stack:
                    self.list_stack.pop()
            elif token_type == "blockquote_open":
                self.blockquote_depth += 1
            elif token_type == "blockquote_close":
                self.blockquote_depth = max(0, self.blockquote_depth - 1)
            elif token_type == "heading_open":
                inline = tokens[index + 1]
                level = max(1, min(int(token.tag[1:]), 6))
                paragraph = self._new_context_paragraph(f"Heading {level}")
                self.render_inline(paragraph, inline.children)
                index += 2
            elif token_type == "paragraph_open":
                inline = tokens[index + 1]
                page_number = self._single_page_marker(inline)
                formula = self._single_display_formula(inline)
                if page_number is not None:
                    from docx.enum.text import WD_BREAK

                    paragraph = self._new_context_paragraph()
                    paragraph.add_run().add_break(WD_BREAK.PAGE)
                elif (
                    formula is not None
                    and not self.list_stack
                    and not self.blockquote_depth
                ):
                    _append_display_formula(
                        self.document,
                        formula.raw,
                        self.summary,
                    )
                else:
                    paragraph = self._new_context_paragraph()
                    self.render_inline(paragraph, inline.children)
                index += 2
            elif token_type in {"fence", "code_block"}:
                self._render_code(token.content)
            elif token_type == "hr":
                paragraph = self._new_context_paragraph()
                _add_horizontal_rule(paragraph)
            elif token_type == "table_open":
                index = self._render_table(tokens, index)
            elif token_type == "html_block":
                self._render_html_fragment(token.content)
            elif (
                token_type == "inline"
                and token.content
                and (index == 0 or tokens[index - 1].type.endswith("_close"))
            ):
                paragraph = self._new_context_paragraph()
                self.render_inline(paragraph, token.children)
            elif (
                token.nesting == 0
                and token.content
                and token_type not in {"inline"}
            ):
                self._new_context_paragraph().add_run(
                    _clean_xml_text(token.content)
                )
            index += 1


def _validate_saved_docx(path: Path) -> None:
    from docx import Document

    with zipfile.ZipFile(path) as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise ValueError(f"DOCX ZIP 校验失败：{bad_member}")
        archive.read("word/document.xml")
    Document(path)


def write_docx(
    markdown: str,
    output_path: Path | str,
    base_dir: Path | str,
) -> ExportSummary:
    """把 Markdown 原子写入 DOCX；失败时保留已有目标文件。"""

    _require_dependencies()

    target = Path(output_path)
    source_directory = Path(base_dir).resolve()
    target.parent.mkdir(parents=True, exist_ok=True)
    temporary = target.with_name(
        f".{target.stem}.{uuid.uuid4().hex}.tmp.docx"
    )
    summary = ExportSummary()
    document = _new_document()
    code_protected, code_regions = _protect_markdown_code(
        _clean_xml_text(str(markdown))
    )
    page_protected, page_markers = _protect_page_markers(code_protected)
    protected, formulas = _protect_formulas(page_protected)
    protected = _restore_markdown_code(protected, code_regions)
    tokens = _markdown_parser().parse(protected)
    _DocxRenderer(
        document,
        base_dir=source_directory,
        formulas=formulas,
        page_markers=page_markers,
        summary=summary,
    ).render(tokens)

    try:
        document.save(temporary)
        _validate_saved_docx(temporary)
        os.replace(temporary, target)
    finally:
        temporary.unlink(missing_ok=True)

    return summary
